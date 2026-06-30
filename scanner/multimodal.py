#!/usr/bin/env python3
"""
multimodal.py -- detect hidden / AI-directed instructions in NON-HTML postings.

The HTML text layer (detect.py) is blind to postings delivered as an image, a PDF, or a DOCX,
and to instructions painted into pixels. This module covers that gap:

  * image  -- OCR the picture and run the AI-instruction matchers over the recovered text;
              a second contrast-stretched OCR pass catches faint / low-contrast "typographic"
              text a human skims past. Text delivered as an image is itself an extraction-evasion
              signal (the agent's OCR sees what a regex JD-scan never would).
  * stego  -- a BEST-EFFORT least-significant-bit heuristic: decode the LSB bitstream and flag
              long printable / instruction-shaped runs. Labeled REVIEW, never CONFIRMED -- LSB
              analysis is noisy and this is a triage hint, not proof.
  * pdf    -- extract text and flag white / near-zero-size / off-page runs, plus document metadata,
              then match instruction language over the hidden portion.
  * docx   -- extract hidden / white / vanish runs and core properties, then match.

All findings use detect.py's (vector, evidence, severity, confidence) shape so they flow through
score()/classify() unchanged. Every dependency is OPTIONAL and imported lazily: if it's missing
the function returns a single informational finding with an install hint and degrades gracefully.

Prior art: Cloud Security Alliance, "Image-based Prompt Injection" (2026) -- vision models cannot
separate visual content from instructions inside it; PhantomLint (arXiv 2508.17884) -- principled
detection of hidden LLM prompts in structured documents.
"""
import os
import re

import detect as DET
import jd_injection_scanner as S


def _is_instruction(text):
    return bool(text) and bool(S.INJECTION_RE.search(text) or DET.SEMANTIC_RE.search(text))


def _missing(dep, hint):
    return [("dependency_missing", f"{dep} not installed -- {hint}", "informational", "low")]


# --------------------------------------------------------------------------- image (OCR)
def _ocr(image):
    import pytesseract
    return pytesseract.image_to_string(image) or ""


def scan_image(path):
    """OCR an image posting + a contrast-stretched pass; flag AI-instruction text and stego."""
    findings = []
    try:
        from PIL import Image, ImageOps
    except ImportError:
        return _missing("Pillow", "pip install Pillow pytesseract && brew install tesseract")
    try:
        img = Image.open(path).convert("RGB")
    except Exception as e:
        return [("image_unreadable", f"{type(e).__name__}: {e}", "informational", "low")]

    # pass 1: as-is OCR. OCR deps are optional -- if absent, note it but still run stego below.
    texts = {}
    ocr_ok = True
    try:
        texts["plain"] = _ocr(img)
    except ImportError:
        ocr_ok = False
        findings.append(("dependency_missing",
                         "pytesseract not installed -- OCR skipped; "
                         "pip install pytesseract && brew install tesseract", "informational", "low"))
    except Exception as e:
        texts["plain"] = ""
        findings.append(("ocr_error", f"{type(e).__name__}: {e}", "informational", "low"))

    # pass 2: autocontrast + grayscale -> surfaces faint/low-contrast text a human skims past
    if ocr_ok:
        try:
            texts["contrast"] = _ocr(ImageOps.autocontrast(ImageOps.grayscale(img), cutoff=1))
        except Exception:
            texts["contrast"] = ""

    plain_norm = re.sub(r"\s+", " ", texts.get("plain", "")).strip()
    contrast_norm = re.sub(r"\s+", " ", texts.get("contrast", "")).strip()

    for label, txt in (("ocr", plain_norm), ("ocr_contrast", contrast_norm)):
        if _is_instruction(txt):
            # an instruction visible only after contrast-stretch is the faint-text canary
            faint = label == "ocr_contrast" and not _is_instruction(plain_norm)
            findings.append(("image_ai_instruction",
                             f"[{'faint/low-contrast ' if faint else ''}OCR] {txt[:200]}",
                             "critical" if faint else "high", "high"))
            break

    findings += _stego_lsb(img)
    return findings or [("image_clean", f"OCR found no instruction text ({len(plain_norm)} chars)",
                         "informational", "low")]


def _stego_lsb(img):
    """Best-effort: read the LSB of each RGB byte, pack to bytes, look for a printable run."""
    try:
        data = list(img.getdata())
    except Exception:
        return []
    bits = []
    for px in data[:200000]:               # cap work; a payload front-loads anyway
        for c in px[:3]:
            bits.append(c & 1)
    out = bytearray()
    for i in range(0, len(bits) - 7, 8):
        byte = 0
        for b in range(8):
            byte = (byte << 1) | bits[i + b]
        out.append(byte)
    try:
        text = out.decode("ascii", "ignore")
    except Exception:
        return []
    # longest run of printable chars
    runs = re.findall(r"[\x20-\x7e]{12,}", text)
    for r in runs:
        if _is_instruction(r):
            return [("image_stego", f"LSB-decoded printable run (best-effort): {r[:160]}",
                     "medium", "low")]
    return []


# --------------------------------------------------------------------------- PDF
def scan_pdf(path):
    """Flag white / near-zero-size / off-page text + metadata, then match instruction language."""
    try:
        from pdfminer.high_level import extract_pages
        from pdfminer.layout import LTTextContainer, LTChar
    except ImportError:
        return _missing("pdfminer.six", "pip install pdfminer.six")
    findings = []
    hidden_chunks = []
    try:
        for page in extract_pages(path):
            pw = getattr(page, "width", 0) or 0
            ph = getattr(page, "height", 0) or 0
            for el in page:
                if not isinstance(el, LTTextContainer):
                    continue
                txt = el.get_text().strip()
                if not txt:
                    continue
                chars = [c for line in el for c in (line if hasattr(line, "__iter__") else [])
                         if isinstance(c, LTChar)]
                size = min((c.size for c in chars), default=12)
                color = _pdf_color(chars)
                x0, y0 = getattr(el, "x0", 0), getattr(el, "y0", 0)
                offpage = (pw and (el.x1 < 0 or el.x0 > pw)) or (ph and (el.y1 < 0 or el.y0 > ph))
                white = color is not None and color >= 0.95
                tiny = size <= 1.0
                if white or tiny or offpage:
                    why = ("white" if white else "") + (" tiny" if tiny else "") + (" offpage" if offpage else "")
                    hidden_chunks.append((why.strip(), txt))
    except Exception as e:
        return [("pdf_error", f"{type(e).__name__}: {e}", "informational", "low")]

    for why, txt in hidden_chunks:
        sev = "critical" if _is_instruction(txt) else "medium"
        vec = "pdf_hidden_instruction" if _is_instruction(txt) else "pdf_hidden_text"
        findings.append((vec, f"[{why}] {txt[:200]}", sev, "high" if sev == "critical" else "medium"))

    findings += _pdf_metadata(path)
    return findings or [("pdf_clean", "no hidden/instruction text found", "informational", "low")]


def _pdf_color(chars):
    """Approximate grayscale brightness 0..1 of a run's fill color, or None if unknown."""
    vals = []
    for c in chars:
        col = getattr(c, "graphicstate", None)
        ncolor = getattr(col, "ncolor", None) if col else None
        if isinstance(ncolor, (int, float)):
            vals.append(float(ncolor))
        elif isinstance(ncolor, (tuple, list)) and ncolor:
            vals.append(sum(float(x) for x in ncolor) / len(ncolor))
    return (sum(vals) / len(vals)) if vals else None


def _pdf_metadata(path):
    try:
        from pdfminer.pdfparser import PDFParser
        from pdfminer.pdfdocument import PDFDocument
    except ImportError:
        return []
    try:
        with open(path, "rb") as fh:
            doc = PDFDocument(PDFParser(fh))
        out = []
        for info in (doc.info or []):
            for k, v in info.items():
                val = v.decode("latin-1", "ignore") if isinstance(v, bytes) else str(v)
                if _is_instruction(val):
                    out.append(("pdf_metadata", f"{k}={val[:160]}", "high", "medium"))
        return out
    except Exception:
        return []


# --------------------------------------------------------------------------- DOCX
def scan_docx(path):
    """Flag hidden / vanish / white runs + core properties, then match instruction language."""
    try:
        import docx  # python-docx
    except ImportError:
        return _missing("python-docx", "pip install python-docx")
    try:
        d = docx.Document(path)
    except Exception as e:
        return [("docx_error", f"{type(e).__name__}: {e}", "informational", "low")]
    findings = []
    for para in d.paragraphs:
        for run in para.runs:
            txt = (run.text or "").strip()
            if not txt:
                continue
            hidden = _docx_run_hidden(run)
            if hidden:
                sev = "critical" if _is_instruction(txt) else "medium"
                vec = "docx_hidden_instruction" if _is_instruction(txt) else "docx_hidden_text"
                findings.append((vec, f"[{hidden}] {txt[:200]}", sev,
                                 "high" if sev == "critical" else "medium"))
    try:
        cp = d.core_properties
        for field in ("comments", "subject", "keywords", "category", "content_status"):
            val = getattr(cp, field, "") or ""
            if _is_instruction(val):
                findings.append(("docx_metadata", f"{field}={val[:160]}", "high", "medium"))
    except Exception:
        pass
    return findings or [("docx_clean", "no hidden/instruction text found", "informational", "low")]


def _docx_run_hidden(run):
    """Return a reason string if a run is vanish/white/tiny, else ''. Reads the run's XML props."""
    try:
        rpr = run._element.rPr
        if rpr is None:
            return ""
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        if rpr.find(f"{ns}vanish") is not None:
            return "vanish"
        color = rpr.find(f"{ns}color")
        if color is not None and (color.get(f"{ns}val") or "").lower() in ("ffffff", "white"):
            return "white"
        sz = rpr.find(f"{ns}sz")
        if sz is not None:
            try:
                if int(sz.get(f"{ns}val")) <= 2:   # half-points; <=2 => <=1pt
                    return "tiny"
            except (TypeError, ValueError):
                pass
        return ""
    except Exception:
        return ""


# --------------------------------------------------------------------------- dispatch
SCANNERS = {"image": scan_image, "pdf": scan_pdf, "docx": scan_docx}
EXT = {".png": "image", ".jpg": "image", ".jpeg": "image", ".bmp": "image", ".gif": "image",
       ".webp": "image", ".tiff": "image", ".pdf": "pdf", ".docx": "docx"}


def scan_file(path, kind=None):
    kind = kind or EXT.get(os.path.splitext(path)[1].lower())
    fn = SCANNERS.get(kind)
    if not fn:
        raise ValueError(f"unsupported file kind for {path!r}; pass kind in {sorted(SCANNERS)}")
    return fn(path)


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("usage: python multimodal.py <image|pdf|docx file>")
        sys.exit(2)
    f = scan_file(sys.argv[1])
    label, top = DET.classify(f)
    print(f"{sys.argv[1]}  score={DET.score(f)} label={label or 'CLEAN'} findings={len(f)}")
    for v, e, sev, conf in f:
        print(f"  [{sev:13} {conf:9}] {v}: {e[:140]!r}")
